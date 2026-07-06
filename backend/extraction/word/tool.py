import os
import uuid
from typing import List, Dict, Any, Optional

from docx import Document
from backend.core.schemas import NormalizedBlock, SourceRef, as_dicts
from backend.extraction.ppt.tool import _detect_image_ext, _save_image_blob


def _safe_text(value) -> str:
    if value is None:
        return ""
    try:
        return str(value).strip()
    except Exception:
        return ""


def _is_heading(paragraph) -> bool:
    try:
        style = (paragraph.style.name or "").lower()
    except Exception:
        return False
    return style.startswith("heading") or style.startswith("title")


class WordExtractorTool:
    name = "word_extraction"

    def run(self, state: dict, config: dict) -> dict:
        file_path = state["file_path"]
        document_id = state.get("document_id")

        if os.path.splitext(file_path)[1].lower() == ".doc":
            state.setdefault("errors", []).append(
                "word_extraction: legacy .doc not supported, convert to .docx"
            )
            state["blocks"] = []
            return state

        blocks = self._extract(file_path, document_id, config, state)
        state["blocks"] = as_dicts(blocks)
        state.setdefault("errors", [])
        return state

    def _detect_language(self, text: str, default: str = "en") -> str:
        if not text or not text.strip():
            return default
        try:
            from langdetect import detect
            return detect(text)
        except Exception:
            return default

    def _extract(self, file_path, document_id, config, state) -> List[NormalizedBlock]:
        blocks = []
        cfg = config or {}
        doc_id = str(document_id) if document_id else str(uuid.uuid4())
        filename = os.path.basename(file_path)

        try:
            doc = Document(file_path)
        except Exception as e:
            state.setdefault("errors", []).append(f"word_extraction: could not open file ({e})")
            return blocks

        blocks += self._paragraphs(doc, doc_id, filename, cfg)
        blocks += self._tables(doc, doc_id, filename, cfg, state)
        blocks += self._images(doc, doc_id, filename, cfg, state)
        return blocks

    def _paragraphs(self, doc, doc_id, filename, cfg) -> List[NormalizedBlock]:
        blocks = []
        for i, para in enumerate(doc.paragraphs):
            text = _safe_text(para.text)
            if not text:
                continue
            try:
                blocks.append(NormalizedBlock(
                    block_id=str(uuid.uuid4()),
                    document_id=doc_id,
                    type="heading" if _is_heading(para) else "text",
                    text=text,
                    source_ref=SourceRef(filename=filename),
                    confidence=cfg.get("extraction_confidence", 1.0),
                    language=self._detect_language(text, cfg.get("default_language", "en")),
                    metadata={
                        "paragraph_index": i,
                        "style": _safe_text(getattr(para.style, "name", "")) or None,
                        "enrichment_failed": cfg.get("enrichment_failed_flag", False),
                    },
                ))
            except Exception:
                continue
        return blocks

    def _tables(self, doc, doc_id, filename, cfg, state) -> List[NormalizedBlock]:
        import pandas as pd
        blocks = []

        for i, table in enumerate(doc.tables):
            try:
                headers = []
                rows = []
                for r, row in enumerate(table.rows):
                    cells = [_safe_text(c.text) for c in row.cells]
                    if r == 0:
                        headers = cells
                    else:
                        rows.append(cells)

                df = pd.DataFrame(rows, columns=headers if headers else None)
                try:
                    md = df.to_markdown(index=False)
                except Exception:
                    md = "\n".join(" | ".join(r) for r in ([headers] + rows[:20]))

                blocks.append(NormalizedBlock(
                    block_id=str(uuid.uuid4()),
                    document_id=doc_id,
                    type="table",
                    text=md,
                    table_data={"headers": headers, "rows": rows},
                    source_ref=SourceRef(filename=filename),
                    confidence=cfg.get("extraction_confidence", 1.0),
                    language=self._detect_language(md, cfg.get("default_language", "en")),
                    metadata={
                        "table_index": i,
                        "enrichment_failed": cfg.get("enrichment_failed_flag", False),
                    },
                ))
            except Exception as e:
                state.setdefault("errors", []).append(f"word_extraction: skipped table {i} ({e})")
                continue
        return blocks

    def _images(self, doc, doc_id, filename, cfg, state) -> List[NormalizedBlock]:
        blocks = []
        out_dir = os.path.join("uploads", "images", doc_id)
        os.makedirs(out_dir, exist_ok=True)

        try:
            rels = doc.part.rels
        except Exception:
            return blocks

        for rel in rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                data = rel.target_part.blob
                if not data:
                    continue
                block_id = str(uuid.uuid4())
                raw_path, preview_path = _save_image_blob(data, out_dir, block_id)

                blocks.append(NormalizedBlock(
                    block_id=block_id,
                    document_id=doc_id,
                    type="image_caption",
                    text="",
                    source_ref=SourceRef(filename=filename),
                    confidence=cfg.get("extraction_confidence", 1.0),
                    language=cfg.get("default_language", "en"),
                    metadata={
                        "raw_image_path": raw_path,
                        "image_path": preview_path or raw_path,
                        "pending_vision": True,
                        "enrichment_failed": False,
                    },
                ))
            except Exception as e:
                state.setdefault("errors", []).append(f"word_extraction: skipped an image ({e})")
                continue

        return blocks


if __name__ == "__main__":
    test_file = "test-data/test.docx"
    state = {"file_path": test_file, "document_id": "doc-003", "errors": []}
    config = {"extraction_confidence": 0.95, "default_language": "en"}

    result = WordExtractorTool().run(state, config)
    blocks = result["blocks"]

    texts = [b for b in blocks if b["type"] == "text"]
    headings = [b for b in blocks if b["type"] == "heading"]
    tables = [b for b in blocks if b["type"] == "table"]
    images = [b for b in blocks if b["type"] == "image_caption"]

    print(f"{len(texts)} text, {len(headings)} heading, {len(tables)} table, {len(images)} image")
    print("errors:", result.get("errors"))